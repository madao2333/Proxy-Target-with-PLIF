# -*- coding: utf-8 -*-
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = next(p for p in ROOT.glob("*.docx") if "模板" in p.name)
OUT = ROOT / "PLIF在脉冲强化学习中的任务相关优势分析_前三章.docx"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_section_layout(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(27.5)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(35.7)
    section.right_margin = Mm(27.7)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(17.5)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size=9)


def restart_page_numbering(section, start=1, fmt=None) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(start))
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    sect_pr.append(pg)


def setup_header_footer(section, header_text=None, roman=False) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in section.header.paragraphs:
        clear_paragraph(p)
    for p in section.footer.paragraphs:
        clear_paragraph(p)

    if header_text:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(header_text)
        set_run_font(run, size=9)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)
    restart_page_numbering(section, 1, "upperRoman" if roman else None)


def set_styles(doc: Document) -> None:
    for section in doc.sections:
        set_section_layout(section)

    for style_name in ("Normal", "正文1"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(12)
            pf = style.paragraph_format
            pf.first_line_indent = Pt(24)
            pf.line_spacing = Pt(20)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_center_title(doc: Document, text: str, size=15, bold=True, style=None) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体" if bold else "宋体", size=size, bold=bold)


def add_body_para(doc: Document, text: str, style="正文1") -> None:
    p = doc.add_paragraph(style=style if style in [s.name for s in doc.styles] else "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=12)


def add_no_indent_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="正文1" if "正文1" in [s.name for s in doc.styles] else "Normal")
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=12, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=12)
    else:
        run = p.add_run(text)
        set_run_font(run, size=12)


def add_heading(doc: Document, level: int, text: str) -> None:
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    p = doc.add_paragraph("  " + text, style=style)
    for run in p.runs:
        set_run_font(run, east_asia="黑体", size=16 if level == 1 else 14, bold=True)


def add_formula(doc: Document, formula: str) -> None:
    p = doc.add_paragraph(style="公式" if "公式" in [s.name for s in doc.styles] else "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(formula)
    set_run_font(run, east_asia="Times New Roman", size=12)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_algorithm_block(doc: Document, caption: str, lines: list[tuple[str, int]]) -> None:
    add_caption(doc, caption)
    add_horizontal_rule(doc)
    for text, indent in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(14 * indent)
        p.paragraph_format.line_spacing = Pt(13.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=10)
    add_horizontal_rule(doc)
    doc.add_paragraph()


def set_cell_text(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


def table_widths(table, widths_twips):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_twips[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_three_line_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def set_cell_bottom_border(cell, size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size=10.5)


def add_table(doc: Document, caption: str, headers, rows, widths_twips) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table_widths(table, widths_twips)
    set_three_line_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_bottom_border(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, align=align)
    doc.add_paragraph()


def add_references(doc: Document) -> None:
    style = "标题一（无序号）" if "标题一（无序号）" in [s.name for s in doc.styles] else "Heading 1"
    p = doc.add_paragraph("参考文献", style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, east_asia="黑体", size=15, bold=True)
    refs = [
        "[1] Fang W, Yu Z, Chen Y, et al. Incorporating learnable membrane time constant to enhance learning of spiking neural networks[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. 2021: 2661-2671.",
        "[2] Zhang D, Zhang T, Jia S, et al. Multi-Sacle Dynamic Coding Improved Spiking Actor Network for Reinforcement Learning[C]//Proceedings of the AAAI Conference on Artificial Intelligence. 2022, 36(1): 59-67.",
        "[3] Chen D, Peng P, Huang T, et al. Fully spiking actor network with intralayer connections for reinforcement learning[J]. IEEE Transactions on Neural Networks and Learning Systems, 2024, 36(2): 2881-2893.",
        "[4] Xu Z, Bu T, Hao Z, et al. Proxy Target: Bridging the Gap Between Discrete Spiking Neural Networks and Continuous Control[J]. arXiv preprint arXiv:2505.24161, 2025.",
        "[5] Fujimoto S, Hoof H, Meger D. Addressing function approximation error in actor-critic methods[C]//International Conference on Machine Learning. PMLR, 2018: 1587-1596.",
        "[6] Tang G, Kumar N, Yoo R, et al. Deep Reinforcement Learning with Population-Coded Spiking Neural Network for Continuous Control[J]. arXiv preprint arXiv:2010.09635, 2020.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="正文1" if "正文1" in [s.name for s in doc.styles] else "Normal")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        r = p.add_run(ref)
        set_run_font(r, size=12)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_center_title(doc, "本科生毕业论文", size=22, bold=True)
    for _ in range(4):
        doc.add_paragraph()
    add_center_title(doc, "题目：PLIF在脉冲强化学习中的任务相关优势分析", size=16, bold=True)
    for _ in range(3):
        doc.add_paragraph()
    fields = ["学    院：", "专    业：", "年    级：", "姓    名：", "学    号：", "指导教师："]
    for field in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(field + " " * 24)
        set_run_font(run, east_asia="宋体", size=14)


def add_declaration(doc: Document) -> None:
    add_center_title(doc, "独创性声明", size=16, bold=True)
    add_body_para(doc, "本人声明：所呈交的毕业设计（论文），是在指导教师指导下独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本毕业设计（论文）中不包含任何他人已经发表或撰写过的研究成果。对本文研究工作做出贡献的个人和集体，均已在文中以明确方式标明。")
    doc.add_paragraph()
    add_no_indent_para(doc, "论文作者签名：")
    add_no_indent_para(doc, "年    月    日")
    doc.add_paragraph()
    add_body_para(doc, "本人声明：本毕业设计（论文）是本人指导学生完成的研究成果，已经审阅过论文的全部内容。")
    doc.add_paragraph()
    add_no_indent_para(doc, "论文指导教师签名：")
    add_no_indent_para(doc, "年    月    日")


def add_abstracts(doc: Document) -> None:
    add_center_title(doc, "摘  要", size=16, bold=True)
    add_body_para(doc, "脉冲神经网络具有事件驱动计算、时序信息处理和低功耗部署潜力，但其离散脉冲发放机制也给连续控制强化学习带来了梯度传播和动作表达方面的困难。本文围绕可学习膜时间常数的PLIF神经元是否能够提升脉冲Actor在连续控制任务中的学习性能展开研究。在Proxy Target与TD3框架下，本文以固定膜电位衰减系数的LIF神经元为主要对照，分析PLIF在不同MuJoCo任务中的性能差异及其可能机制。现有实验结果表明，PLIF在Walker2d-v4任务上表现出稳定优势，五个随机种子均优于LIF，平均最大评估回报提升约48.14%；而在InvertedDoublePendulum-v4中二者基本持平，在Ant-v4、HalfCheetah-v4和Hopper-v4中暂未体现稳定优势。基于相关文献与现有实验，本文将PLIF的作用定位为任务相关的时间尺度自适应能力，而非对LIF的无条件替代。后续章节将从tau演化、固定tau消融以及STBP梯度传播角度进一步解释PLIF在脉冲强化学习中的适用边界。")
    doc.add_paragraph()
    add_no_indent_para(doc, "关键词：脉冲神经网络；强化学习；PLIF神经元；Proxy Target；连续控制；TD3", bold_prefix="关键词：")
    doc.add_page_break()
    add_center_title(doc, "ABSTRACT", size=16, bold=True)
    add_body_para(doc, "Spiking neural networks provide event-driven computation, temporal information processing, and potential energy-efficient deployment, but their discrete spike dynamics make continuous-control reinforcement learning difficult. This thesis studies whether PLIF neurons with learnable membrane time constants can improve the learning performance of spiking actors. Based on the Proxy Target and TD3 framework, fixed-decay LIF neurons are used as the main baseline, and the task-dependent effects of PLIF are analyzed on MuJoCo continuous-control benchmarks. Existing results show that PLIF achieves a stable advantage on Walker2d-v4, outperforming LIF under all five random seeds and improving the average maximum evaluation reward by about 48.14%. In contrast, PLIF is close to LIF on InvertedDoublePendulum-v4 and does not yet show stable gains on Ant-v4, HalfCheetah-v4, or Hopper-v4. Therefore, the thesis interprets PLIF as a task-adaptive temporal-scale modeling mechanism rather than an unconditional replacement for LIF. Subsequent analysis will further examine tau evolution, fixed-tau ablation, and STBP gradient propagation.")
    doc.add_paragraph()
    add_no_indent_para(doc, "KEY WORDS: Spiking neural networks; Reinforcement learning; PLIF neuron; Proxy Target; Continuous control; TD3", bold_prefix="KEY WORDS:")


def add_toc(doc: Document) -> None:
    add_center_title(doc, "目  录", size=16, bold=True)
    entries = [
        ("摘  要", "Ⅰ", 0),
        ("ABSTRACT", "Ⅱ", 0),
        ("第一章  绪论", "1", 0),
        ("1.1  研究背景", "1", 1),
        ("1.2  问题提出", "2", 1),
        ("1.3  研究目标", "2", 1),
        ("1.4  主要贡献", "3", 1),
        ("第二章  相关理论与方法基础", "3", 0),
        ("2.1  强化学习与连续控制", "3", 1),
        ("2.2  脉冲神经网络", "4", 1),
        ("2.3  PLIF神经元", "5", 1),
        ("2.4  代理梯度与STBP", "6", 1),
        ("2.5  Proxy Target框架", "6", 1),
        ("第三章  PLIF脉冲Actor设计", "7", 0),
        ("3.1  整体网络结构", "7", 1),
        ("3.2  LIF与PLIF的计算差异", "7", 1),
        ("3.3  PLIF tau参数化", "8", 1),
        ("3.4  训练流程", "9", 1),
        ("参考文献", "10", 0),
    ]
    for title, page, level in entries:
        p = doc.add_paragraph(style="toc 2" if level else "toc 1")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(24 * level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"{title}\t{page}")
        set_run_font(run, size=12)


def add_chapter_1(doc: Document) -> None:
    add_heading(doc, 1, "绪论")
    add_heading(doc, 2, "研究背景")
    for text in [
        "强化学习研究智能体如何在与环境的连续交互中根据奖励信号改进决策策略。对于机器人控制、自动驾驶、无人系统和具身智能等问题，环境状态和动作空间往往是连续且高维的，策略网络需要在复杂动力学中学习稳定、平滑并具有泛化能力的动作输出。深度强化学习通过将深度神经网络作为函数逼近器，显著提高了连续控制任务中的策略表示能力，但其训练和推理通常依赖密集矩阵运算，能耗和部署成本较高，限制了其在边缘设备和低功耗机器人平台中的应用。",
        "脉冲神经网络被认为是更接近生物神经系统的第三代神经网络模型。与传统人工神经网络使用连续激活值不同，SNN以离散脉冲事件传递信息，神经元内部还显式保留膜电位、电流和阈值等动态状态。这种事件驱动机制使SNN在神经形态芯片上具有低功耗、低延迟和稀疏计算的潜力。Fang等提出的PLIF研究指出，SNN不仅具有时序信息处理能力和生物可解释性，而且膜相关参数会直接影响单个神经元的动态响应与网络表达能力[1]。Tang等的PopSAN工作进一步表明，将population coding引入脉冲Actor后，SNN能够被用于连续控制任务，并可在神经形态硬件上获得明显能效优势[6]。",
        "然而，将SNN直接用于强化学习中的连续控制并不容易。一方面，脉冲发放函数通常由阶跃函数描述，前向传播是离散的，标准反向传播难以直接求导；另一方面，连续控制任务要求Actor输出连续动作，而SNN天然输出的是脉冲序列或脉冲率，二者之间存在表达形式上的差距。已有研究围绕这一问题提出了多种改进方向：Zhang等通过多尺度动态编码增强脉冲Actor的时空表征能力[2]；Chen等尝试构建全脉冲Actor，并指出以脉冲率再接全连接层表示连续动作会引入额外浮点运算，从而削弱全脉冲部署优势[3]；Xu等则从强化学习算法结构出发，提出Proxy Target框架以缓解SNN离散动态与连续控制目标网络更新机制之间的不匹配[4]。",
        "在上述背景下，本文关注PLIF神经元在脉冲强化学习中的作用。PLIF的核心思想是将膜时间常数或膜电位保留系数设为可学习参数，使神经元可以根据任务和训练过程调节自身时间尺度。对于连续控制任务，环境动力学通常包含不同长度的时序依赖，固定时间常数的LIF神经元可能难以同时适应短时响应和长时状态保持需求。因此，研究PLIF是否能在Proxy Target框架下提高SNN Actor的学习性能，既具有神经元动力学层面的理论意义，也具有低功耗连续控制部署方面的实际价值。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "问题提出")
    for text in [
        "传统LIF神经元使用固定的膜电位衰减系数。固定系数简化了模型和训练过程，但也意味着所有任务、所有网络层甚至所有训练阶段共享相同的时间尺度假设。当环境状态变化较快时，较大的膜电位保留可能造成历史信息干扰；当任务依赖更长时间的运动协调时，较小的保留系数又可能使关键时序信息过早衰减。连续控制任务之间存在显著动力学差异，例如倒立摆任务的回报较容易接近饱和，而Walker2d-v4涉及更复杂的双腿协调和接触动力学，对历史状态保持的需求可能更强。",
        "PLIF为这一问题提供了一个可学习的建模入口。Fang等的研究表明，在监督学习任务中共同学习突触权重和膜时间常数可以提升SNN对初值的鲁棒性并加快学习[1]。但在强化学习场景中，策略学习的目标并非固定标签，而是由奖励、价值估计和环境交互共同决定；Actor还要与Critic、Replay Buffer、目标网络或Proxy Target形成耦合。因此，PLIF在分类任务中的优势不能直接推断到连续控制强化学习中，需要结合具体算法框架重新验证。",
        "此外，本文实验框架中同时存在两个名称相同但含义不同的tau。第一个tau来自PLIF神经元，表示膜电位衰减或保留相关的时间尺度参数；第二个tau来自TD3算法，表示目标网络软更新系数。Fujimoto等提出的TD3通过双Critic、延迟策略更新和目标策略平滑缓解Actor-Critic中的函数逼近误差与过估计问题[5]，其中软更新tau服务于目标网络稳定性。若不区分这两个tau，就容易把神经元动态的可学习时间尺度与强化学习算法的目标网络更新机制混为一谈。",
        "因此，本文的问题可以概括为：在Proxy Target + TD3框架下，将SNN Actor中的LIF神经元替换为PLIF神经元，是否能够稳定提升连续控制性能；如果性能提升只在部分任务上出现，这种差异是否能够由任务时序依赖、tau学习趋势和梯度传播状态解释。围绕这一问题，论文不预设PLIF必然全面优于LIF，而是把重点放在任务相关优势和适用边界的分析上。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "研究目标")
    for text in [
        "本文的总体目标是在已有Proxy Target代码与实验结果基础上，系统分析PLIF神经元在脉冲强化学习连续控制任务中的作用。具体而言，首先在相同Actor-Critic训练框架、相同环境和相同随机种子设置下，对比LIF与PLIF两类脉冲Actor的最大评估回报、学习曲线和seed级别胜负关系，判断PLIF优势是否稳定存在。",
        "其次，本文拟分析PLIF中tau参数在不同任务和不同网络层中的演化趋势。若PLIF在某一任务中显著优于LIF，则需要进一步观察隐藏层和输出层tau是否偏离固定LIF值，以及这种偏离是否与更强历史信息保持或更快动态响应相一致。对于PLIF没有明显优势的任务，也需要分析tau是否学习不足、是否收敛到接近LIF的固定值，或是否可能增加优化难度。",
        "再次，本文将通过固定tau LIF消融实验检验PLIF优势来源。如果某个固定tau设置能够接近PLIF性能，则说明PLIF主要起到了自动选择合适时间尺度的作用；如果PLIF仍优于所有固定tau设置，则说明按层学习、训练过程中的动态调节或与梯度传播的交互可能更加重要。最后，本文将结合STBP梯度打印结果，分析PLIF对脉冲Actor训练稳定性、梯度范数和tau参数更新的影响。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "主要贡献")
    for text in [
        "本文的第一项贡献是将PLIF神经元引入Proxy Target + TD3脉冲强化学习框架，并在多个MuJoCo连续控制环境上与固定tau的LIF神经元进行对齐比较。现有结果显示，PLIF在Walker2d-v4上取得稳定且显著的优势，五个随机种子均优于LIF，平均最大评估回报提升约48.14%；在InvertedDoublePendulum-v4上两者基本持平，而在Ant-v4、HalfCheetah-v4和Hopper-v4中暂未表现出稳定优势。",
        "第二项贡献是将PLIF的性能结论从“整体优劣比较”推进到“任务相关机制分析”。本文不把PLIF表述为LIF的全面替代，而是强调可学习tau为SNN Actor提供了任务自适应的时间尺度建模能力。这一表述既与PLIF文献中膜时间常数影响神经元动态的结论一致[1]，也与连续控制任务中不同环境动力学需求存在差异的实际情况相吻合。",
        "第三项贡献是明确区分PLIF神经元tau与TD3软更新tau，并结合代码实现说明本文实验中的PLIF tau更接近膜电位衰减或保留系数。该澄清有助于避免概念混淆，也为后续固定tau消融、tau学习率设置和STBP梯度分析提供清晰的变量定义。",
    ]:
        add_body_para(doc, text)


def add_chapter_2(doc: Document) -> None:
    add_heading(doc, 1, "相关理论与方法基础")
    add_heading(doc, 2, "强化学习与连续控制")
    for text in [
        "强化学习通常用马尔可夫决策过程描述。智能体在时刻t观察环境状态s_t，选择动作a_t，环境返回奖励r_t并转移到下一状态s_{t+1}。策略学习的目标是在长期交互中最大化折扣累计回报。与离散动作任务相比，连续控制任务要求策略直接输出连续动作向量，动作维度和状态维度通常较高，且系统动力学对动作扰动较敏感。因此，连续控制算法需要同时保证价值估计稳定性和策略更新平滑性。",
        "Actor-Critic框架将策略函数与价值函数分开建模。Actor根据当前状态输出动作，Critic估计状态-动作对的价值，并为Actor提供优化方向。确定性策略梯度方法适合连续动作空间，但在使用神经网络作为函数逼近器时，价值估计误差会被策略更新放大。Fujimoto等指出，Actor-Critic中同样存在过估计偏差和误差累积问题，并提出TD3算法进行改进[5]。",
        "TD3的核心机制包括三点。第一，使用两个Critic网络同时估计Q值，并在构造目标值时取二者较小值，以降低过估计风险。第二，延迟策略更新，即Critic更新多次后再更新Actor，从而减少Actor追随不稳定价值估计的频率。第三，目标策略平滑，即在目标动作上加入裁剪噪声，使Critic学习到对相近动作更平滑的价值估计。本文实验以TD3为连续控制基础算法，因此后续所有SNN Actor的比较都建立在这一Actor-Critic框架之上。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "脉冲神经网络")
    for text in [
        "脉冲神经网络以离散脉冲作为主要信息载体。神经元在每个仿真时间步接收突触输入，更新膜电位，当膜电位超过阈值时发放脉冲并执行重置。与ANN的连续激活不同，SNN的状态演化具有显式时间维度，能够在膜电位和脉冲序列中保存历史输入信息。正因为这一特点，SNN被广泛用于低功耗、事件驱动和时序信息处理场景。",
        "LIF神经元是最常用的SNN基本模型之一。其膜电位在未发放脉冲时按固定比例保留历史状态，并叠加当前输入；当膜电位超过阈值后，神经元输出脉冲并重置或削弱膜电位。该模型结构简单、计算成本低，但固定膜电位衰减系数也限制了神经元对不同时间尺度的适应能力。Fang等指出，膜相关参数决定单个神经元的动态特性，如果所有神经元共享同一固定时间常数，会削弱神经元异质性和网络表达能力[1]。",
        "在强化学习连续控制中，SNN还需要解决动作表达问题。Tang等提出的PopSAN采用population coding将连续状态编码为脉冲群体活动，再通过脉冲Actor与深度Critic结合完成连续控制[6]。Zhang等提出的MDC-SAN进一步将网络尺度的群体编码与神经元尺度的二阶动态编码结合，以增强空间-时间状态表征[2]。Chen等则指出，许多方法使用脉冲率再经过全连接层表示连续动作，这会引入浮点矩阵运算；他们提出使用非脉冲神经元膜电位和层内连接构造更完整的全脉冲Actor[3]。这些研究共同说明，SNN用于连续控制的关键并不仅是替换网络结构，还包括编码、解码、神经元动态和训练算法的协同设计。",
    ]:
        add_body_para(doc, text)

    doc.add_page_break()
    add_table(
        doc,
        "表2-1  相关脉冲强化学习与PLIF研究对本文的启示",
        ["文献", "核心思想", "对本文的启示"],
        [
            ("Fang等[1]", "学习膜时间常数，增强SNN神经元异质性与学习能力。", "为PLIF tau可学习提供直接理论依据。"),
            ("Tang等[6]", "将population coding用于连续控制脉冲Actor，并验证神经形态部署能效。", "支撑本文采用群体编码/解码结构。"),
            ("Zhang等[2]", "多尺度动态编码增强SNN Actor的时空表征。", "说明连续控制任务需要更强时序动态建模。"),
            ("Chen等[3]", "构建带层内连接的全脉冲Actor，减少非脉冲浮点运算。", "提示动作解码方式会影响SNN部署优势。"),
            ("Xu等[4]", "以Proxy Target缓解SNN离散动态与连续控制软更新机制的矛盾。", "构成本文实验框架。"),
        ],
        [1800, 3780, 3780],
    )

    add_heading(doc, 2, "PLIF神经元")
    for text in [
        "PLIF神经元可以看作对LIF神经元的参数化扩展。LIF中的膜电位保留或衰减系数通常由人工设定；PLIF则将该参数纳入可学习变量，使其能够随误差反向传播更新。Fang等将这种模型称为Parametric Leaky Integrate-and-Fire神经元，并指出可学习膜时间常数能够使网络对初始值不那么敏感，同时加快学习过程[1]。",
        "需要注意的是，不同实现中tau的数学含义可能存在差异。在神经元动力学文献中，tau常被称为膜时间常数，用于描述膜电位随时间衰减的速度；在本文代码中，PLIFNode通过sigmoid(w)得到一个位于0到1之间的系数，并在更新中写作膜电位的保留系数。由于当前LIF固定值为0.75，PLIF初始值也与该保留系数对齐。因此，本文后续讨论tau时，主要指代码实现中的膜电位保留/衰减系数，而不是连续时间微分方程中的物理时间常数。",
        "从连续控制角度看，PLIF的意义在于允许不同网络层学习不同的时间尺度。隐藏层tau较大时，膜电位对历史输入保留更强，有助于建模较长时间依赖；tau较小时，神经元响应更偏向当前输入，有助于快速跟随状态变化。输出population层的tau还会影响动作脉冲活动的平滑程度。因此，PLIF是否有效不只取决于是否可学习，还取决于任务动力学、层级功能和强化学习损失共同塑造出的tau演化路径。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "代理梯度与STBP")
    for text in [
        "SNN训练的主要困难来自脉冲发放函数不可导。阶跃函数在阈值处不连续，在其他位置梯度几乎为零，直接使用反向传播会导致参数难以更新。代理梯度方法用一个连续函数的导数近似阶跃函数的梯度，使网络在前向传播中保持脉冲离散性，在反向传播中获得可用梯度。Chen等指出，代理梯度已经被用于训练多层SNN，并使其在连续控制任务中能够达到接近对应深度网络的性能[3]。",
        "STBP将反向传播扩展到SNN的时间维度和层级维度。对于包含T个仿真时间步的脉冲网络，损失函数不仅通过网络层反传，也沿神经元状态的时间递推关系反传。因此，膜电位保留系数、脉冲发放稀疏性和重置机制都会影响梯度大小和传播路径。在本文实验中，STBP梯度分析用于观察不同层、不同时间步和不同神经元类型下的梯度范数，从而判断PLIF是否改善了梯度过小、传播不稳定或脉冲活动过稀疏等问题。",
        "对于PLIF而言，STBP还承担了更新tau参数的功能。tau通过膜电位递推参与每个时间步的状态更新，其梯度来自Actor损失对膜电位动态的间接影响。如果tau参数梯度长期接近零，说明PLIF虽形式上可学习，但实际训练中可能难以调整时间尺度；如果tau梯度波动过大，则可能带来优化不稳定。因此，tau梯度和普通权重梯度应作为后续机制分析的重要证据。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "Proxy Target框架")
    for text in [
        "传统TD3使用目标Actor网络产生下一状态动作，并通过软更新使目标网络缓慢跟随当前Actor。该机制在ANN中通常稳定有效，但SNN Actor的离散脉冲动态和不可导发放函数会使软更新后的目标网络表现出不连续或不平滑的行为。Xu等指出，连续控制算法中的目标网络软更新机制与SNN的离散非可导动态存在冲突，可能造成训练不稳定和性能下降[4]。",
        "Proxy Target框架的核心思想是引入一个连续、可微的代理目标网络。训练时，当前SNN Actor输出动作，Proxy Target ANN通过均方误差拟合该动作输出；Critic构造目标值时使用Proxy Target产生下一状态动作。由于Proxy Target具有连续可微动态，其输出更适合TD3中的目标动作平滑和价值目标构造。与此同时，Proxy Target只在训练阶段使用，部署时仍保留SNN Actor，因此不会增加推理阶段的能耗开销[4]。",
        "本文实验基于Proxy Target + TD3展开。代码中在Proxy Target更新阶段对SNN Actor输出使用torch.no_grad()，使Actor动作只作为监督标签，不让proxy loss反向传播回SNN Actor。这样既符合Proxy Target的优化目标，也避免每次proxy iteration都额外触发一次SNN Actor的STBP反向计算。随后算法按TD3流程更新Critic，并以延迟策略更新方式优化SNN Actor。对于PLIF Actor，tau参数与普通Actor权重采用不同学习率，以降低时间尺度参数过快变化造成的不稳定风险。",
    ]:
        add_body_para(doc, text)


def add_chapter_3(doc: Document) -> None:
    add_heading(doc, 1, "PLIF脉冲Actor设计")
    add_heading(doc, 2, "整体网络结构")
    for text in [
        "本文的脉冲Actor采用“Population Spike Encoder—Spike MLP—Population Decoder”的整体结构。该结构继承了PopSAN类方法的基本思想：先用群体神经元把连续状态编码为脉冲序列，再由脉冲多层感知机在若干仿真时间步内完成时序计算，最后把输出population活动解码为连续动作。这样的设计既保留了SNN事件驱动和时序计算的特点，又能够满足连续控制任务对动作向量的输出要求。",
        "输入端的Population Spike Encoder将每个状态维度扩展为若干个编码神经元。每个编码神经元对应不同的感受野均值，输入状态经过高斯形状的population活动函数后，在多个仿真时间步中生成规则脉冲序列。与单神经元编码相比，population coding提高了连续状态的表示分辨率，也为后续SNN层提供更丰富的稀疏输入。Tang等的研究表明，population-coded spiking actor能够与深度Critic结合用于连续控制任务[6]，这为本文结构选择提供了直接参考。",
        "中间的Spike MLP由若干线性突触层和脉冲神经元层组成。对于LIF、CLIF和PLIF三种设置，线性层结构保持一致，差异主要体现在神经元状态更新方式上。隐藏层负责从输入脉冲序列中提取时序特征，输出population层则累计各时间步的脉冲活动，并在仿真结束后取平均作为动作解码输入。",
        "输出端的Population Decoder按照动作维度分组，将每个动作维度对应的一组population活动映射为一个连续动作分量，并通过tanh函数和动作范围系数限制输出。相比直接使用单个脉冲神经元表示动作，population decoder能够利用一组神经元的活动分布表达连续值，从而缓解脉冲离散性与连续动作空间之间的矛盾。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "LIF与PLIF的计算差异")
    for text in [
        "在当前实现中，LIF和PLIF共享相同的突触输入计算方式，二者的差异主要体现在膜电位保留系数是否可学习。为避免与TD3目标网络软更新系数τ_TD3混淆，本文在本节中用α表示神经元膜电位保留系数。设t为SNN仿真时间步，k为网络层，I_t^k为第k层在t时刻的突触输入，V_t^k为膜电位，S_t^k为脉冲输出，Θ为阶跃发放函数，σ为sigmoid函数，则第k层神经元状态递推可写为：",
    ]:
        add_body_para(doc, text)
    add_formula(doc, "I_t^k = W^k S_t^(k-1) + b^k    (3-1)")
    add_formula(doc, "V_t^k = α^k V_(t-1)^k (1 - S_(t-1)^k) + I_t^k    (3-2)")
    add_formula(doc, "S_t^k = Θ(V_t^k - V_th)    (3-3)")
    add_body_para(doc, "其中(1 - S_(t-1)^k)表示上一时间步发放脉冲后对膜电位进行重置或抑制。该形式使发放过脉冲的神经元在下一步不会简单累积过高历史膜电位，从而维持脉冲发放的离散事件特性。")
    add_body_para(doc, "在上述统一记号下，LIF与PLIF的区别可以进一步写成膜电位保留系数α^k的不同取值。LIF使用固定保留系数α_LIF=0.75；PLIF则将该系数参数化为α_PLIF^k=σ(w^k)，其中w^k为第k层PLIF神经元模块中的可学习参数。")
    add_formula(doc, "α^k = α_LIF = 0.75    (LIF)    (3-4)")
    add_formula(doc, "α^k = α_PLIF^k = σ(w^k),    w^k learnable    (PLIF)    (3-5)")
    for text in [
        "因此，PLIF并未改变Actor的输入输出接口，也未改变突触层数量和隐藏单元规模，而是在神经元状态递推中引入可学习时间尺度。若α_PLIF^k学习到较大值，膜电位保留更强，历史输入对当前状态影响更久；若学习到较小值，神经元更快遗忘历史状态，对当前输入更加敏感。",
        "这种差异看似只改变一个系数，但在多时间步STBP中会影响整个梯度传播路径。由于膜电位递推跨越多个仿真时间步，α_PLIF不仅影响前向脉冲活动，也影响损失函数对早期输入和早期膜电位状态的梯度。代码日志中输出的tau对应本节的α_PLIF，而不是TD3中的τ_TD3。PLIF的潜在优势正来自这种前向动态与反向梯度路径的共同调整。",
    ]:
        add_body_para(doc, text)

    add_table(
        doc,
        "表3-1  LIF与PLIF神经元状态更新差异",
        ["神经元类型", "膜电位保留系数", "参数更新", "本文关注点"],
        [
            ("LIF", "α_LIF=0.75", "仅更新突触权重", "作为固定时间尺度基线。"),
            ("PLIF", "α_PLIF^k=σ(w^k)，初始对齐0.75", "同时更新突触权重与w^k", "分析任务相关时间尺度自适应。"),
            ("CLIF", "固定电压衰减并引入电流衰减", "更新突触权重", "作为补充对照，区分电流动态影响。"),
        ],
        [1450, 2450, 2450, 3010],
    )

    add_heading(doc, 2, "PLIF tau参数化")
    for text in [
        "当前代码中的PLIFNode使用一个可学习标量w表示每个PLIF神经元模块的时间尺度参数，并通过sigmoid函数将其映射到0到1之间。初始化时，w根据目标初始tau反推得到，使PLIF的初始膜电位保留系数与LIF固定值0.75对齐。这样可以保证LIF与PLIF比较时，初始动态尽可能一致，避免把初值差异误判为可学习机制带来的性能变化。",
        "在Spike MLP中，每个隐藏层和输出population层各配置一个PLIFNode。因此，在两层隐藏层设置下，PLIF Actor通常包含hidden0、hidden1和output三个tau轨迹。按层设置tau比全网络共享一个tau更灵活，因为不同层承担的功能并不相同：浅层更接近编码输入，可能需要快速响应状态变化；深层负责组合时序特征，可能需要更长历史保持；输出层直接影响动作脉冲活动，其tau变化会影响动作平滑性。",
        "tau参数与普通权重使用不同学习率。代码中PLIF参数单独分组，Proxy Target版本默认将PLIF学习率设置为proxy_lr的一定比例。这样做的原因是tau直接控制神经元状态递推，变化过快可能导致脉冲发放率、膜电位分布和梯度范数同时波动，从而破坏强化学习训练稳定性。后续实验如果观察到tau变化过慢或过快，需要结合PLIF学习率进一步分析。",
        "本文在解释tau曲线时，需要始终围绕“膜电位保留系数”这一代码语义展开。若某任务中hidden tau明显升高，可以解释为模型倾向于保留更长历史信息；若output tau降低，则可能表示输出动作更依赖当前时刻特征或需要更快更新。该解释不能简单等同于TD3目标网络软更新tau，也不能直接套用连续时间神经元模型中的物理时间常数含义。",
    ]:
        add_body_para(doc, text)

    add_heading(doc, 2, "训练流程")
    for text in [
        "本文训练流程以Proxy Target + TD3为主线。Proxy Actor仅在训练阶段作为连续、可微的目标策略近似器使用；部署阶段仍保留SNN Actor。为清晰说明各网络之间的更新关系，算法3-1给出本文采用的训练流程，其中τ_TD3表示目标Critic软更新系数，PLIF神经元中的膜电位保留系数仍记为α_PLIF。",
    ]:
        add_body_para(doc, text)
    add_algorithm_block(
        doc,
        "算法3-1  Proxy Target + TD3下的PLIF脉冲Actor训练流程",
        [
            ("Input: 环境E，Replay Buffer D，SNN Actor π_φ^SNN，Proxy Actor π_φ′^P，", 0),
            ("双Critic Q_θ1, Q_θ2，目标Critic Q_θ1′, Q_θ2′，", 1),
            ("折扣因子γ，目标网络软更新系数τ_TD3，proxy更新次数K，batch size N，延迟策略更新频率d", 1),
            ("Output: 训练后的SNN Actor π_φ^SNN", 0),
            ("Initialize π_φ^SNN, π_φ′^P, Q_θ1, Q_θ2, Q_θ1′, Q_θ2′ and D", 0),
            ("for each environment step do", 0),
            ("根据当前SNN Actor选择动作a_t，并与环境交互得到(s_t, a_t, r_t, s_{t+1})", 1),
            ("将转移样本存入D", 1),
            ("for k = 1 to K do", 1),
            ("从D中采样小批量状态s_i", 2),
            ("在no_grad模式下计算动作标签a_i^SNN = π_φ^SNN(s_i)", 2),
            ("更新Proxy Actor，使其最小化L_proxy = 1/N sum_i ||π_φ′^P(s_i) - a_i^SNN||_2^2", 2),
            ("end for", 1),
            ("从D中采样小批量转移(s_i, a_i, r_i, s_i')", 1),
            ("使用Proxy Actor生成目标动作，并加入TD3目标策略平滑噪声", 1),
            ("y_i = r_i + γ min_j Q_θj′(s_i′, a_i′)", 1),
            ("更新双Critic，使其最小化TD误差", 1),
            ("if 当前迭代满足延迟策略更新条件 then", 1),
            ("更新SNN Actor，使L_actor = -1/N sum_i Q_θ1(s_i, π_φ^SNN(s_i))最小", 2),
            ("若Actor为PLIF，同步通过STBP更新突触权重和w^k", 2),
            ("软更新目标Critic: θ_j′ <- τ_TD3 θ_j + (1 - τ_TD3) θ_j′", 2),
            ("end if", 1),
            ("end for", 0),
        ],
    )
    for text in [
        "算法3-1中，Proxy Actor的监督标签由当前SNN Actor在无梯度模式下产生，因此proxy loss只更新Proxy Actor参数，不会反向传播到SNN Actor。Critic更新时使用Proxy Actor产生下一状态动作，并通过双Q较小值构造TD目标；Actor则按照TD3的延迟策略更新机制，由Critic给出的Q值优化当前SNN策略。",
        "当Actor为PLIF脉冲网络时，Actor反向传播会同时更新普通突触权重和w^k参数。由于w^k通过α_PLIF^k=σ(w^k)参与每个仿真时间步的膜电位递推，其梯度需要通过STBP在时间维度上累计。因此，PLIF训练不仅是参数数量增加，也改变了Actor内部时序动态的可调方式。本文后续实验分析将围绕性能证据、α_PLIF轨迹和梯度证据展开，以判断PLIF在Walker2d-v4上的优势是否来自更合适的时间尺度建模，以及为何这种优势没有在所有任务上同时出现。",
    ]:
        add_body_para(doc, text)


def build() -> None:
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    set_styles(doc)
    set_section_layout(doc.sections[0])

    # Cover section: no running header/footer.
    for p in doc.sections[0].header.paragraphs:
        clear_paragraph(p)
    for p in doc.sections[0].footer.paragraphs:
        clear_paragraph(p)
    add_cover(doc)

    declaration_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(declaration_section)
    declaration_section.header.is_linked_to_previous = False
    declaration_section.footer.is_linked_to_previous = False
    for p in declaration_section.header.paragraphs:
        clear_paragraph(p)
    for p in declaration_section.footer.paragraphs:
        clear_paragraph(p)
    add_declaration(doc)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(front)
    setup_header_footer(front, roman=True)
    add_abstracts(doc)
    doc.add_page_break()
    add_toc(doc)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(body)
    setup_header_footer(body, "天津大学2026届本科生毕业设计（论文）", roman=False)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_references(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
